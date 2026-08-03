"""
Learning materials: upload handling + the one-time LLM read (PR 2 of
the walkthrough arc).

A file upload is worth exactly two artifacts: the extracted text and
a one-time digest — Theo's own reading of what the material contains.
The original bytes are deliberately NOT kept (the husband's material
is a law-firm document; we store what the coach needs and nothing
more). Links get no digest call at all: latent knowledge stands in at
prompt time, and the user's walkthrough is the primary source for
every kind anyway.
"""

import io
import os

import anthropic

import db

MODEL = os.environ.get("MATERIAL_MODEL", "claude-sonnet-4-5")

ALLOWED_EXTENSIONS = (".pdf", ".docx")
MAX_FILE_BYTES = 20 * 1024 * 1024
# The digest call reads at most this much of the extracted text; the
# stored extracted_text keeps everything (question generation later
# reads specific chunks, not the whole thing).
DIGEST_INPUT_CHARS = 80_000

DIGEST_PROMPT = """\
You are reading a document a learner has just shared with their
coach. Read it once, carefully, and produce the coach's working notes
on it — in English, under 300 words:

1. What this document IS (one sentence).
2. Its structure: the sections/parts as the author organized them,
   with a rough size for each (items, pages' worth, or paragraphs).
3. What kind of knowledge it holds (facts to recall? procedures?
   concepts? worked examples?) and roughly how many discrete items.
4. Anything that looks central or unusually dense — the parts an
   expert would call the heart of it.

Do NOT evaluate quality, do NOT suggest study strategies, do NOT
address the learner. These notes are for the coach's eyes; the
learner's own account of what matters (collected separately) always
outranks this reading."""


def extract_text(filename, data):
    """→ (text, error). Extracts what the coach can read from an
    uploaded file. Never raises."""
    name = (filename or "").lower()
    try:
        if name.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(data))
            pages = [(page.extract_text() or "") for page in reader.pages]
            text = "\n\n".join(p.strip() for p in pages if p.strip())
        elif name.endswith(".docx"):
            from docx import Document
            doc = Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells
                             if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts)
        else:
            return "", f"unsupported file type: {filename}"
    except Exception as e:
        return "", f"could not read {filename}: {e}"
    if not text.strip():
        return "", (f"{filename} contained no extractable text "
                    f"(a scanned/image PDF?)")
    return text, ""


def register_upload(user_id, filename, data):
    """Validate + extract + register an uploaded file. → (material_id,
    error). The digest is NOT generated here — call digest_material
    afterwards (handlers run it in the background so the upload
    response stays instant)."""
    if not filename or not filename.lower().endswith(ALLOWED_EXTENSIONS):
        return None, "Only PDF and DOCX files are supported."
    if len(data) > MAX_FILE_BYTES:
        return None, "File is larger than 20MB."
    if not data:
        return None, "The file arrived empty."
    text, err = extract_text(filename, data)
    if err:
        return None, err
    material_id = db.add_user_material(
        user_id, "file", title=filename, orig_filename=filename,
        extracted_text=text)
    return material_id, ""


def register_link(user_id, url, title=""):
    """Register a link the user studies from. → (material_id, error)."""
    url = (url or "").strip()
    if not url.startswith(("http://", "https://")):
        return None, "Please paste a full link (starting with http)."
    if len(url) > 2000:
        return None, "That link is too long."
    return db.add_user_material(
        user_id, "link", title=(title or "").strip()[:200] or url,
        source_url=url), ""


def digest_material(material_id):
    """The one-time LLM read of an uploaded file. Stores the digest,
    flight-records the call, emits material_digested. Returns the
    digest text or None. Never raises — a failed digest leaves the
    material usable (the walkthrough is the primary source)."""
    try:
        m = db.get_material(material_id)
        if not m or not (m.get("extracted_text") or "").strip():
            return None
        body = m["extracted_text"][:DIGEST_INPUT_CHARS]
        truncated = len(m["extracted_text"]) > DIGEST_INPUT_CHARS
        messages = [{"role": "user",
                     "content": (f"Document: {m.get('title') or '(untitled)'}"
                                 + (" (truncated for this read)"
                                    if truncated else "")
                                 + f"\n\n{body}")}]
        client = anthropic.Anthropic()
        resp = client.messages.create(
            model=MODEL, max_tokens=800, system=DIGEST_PROMPT,
            messages=messages)
        digest = "".join(b.text for b in resp.content
                         if getattr(b, "type", "") == "text").strip()
        if not digest:
            return None
        call_id = db.save_llm_call(
            m["user_id"], "material_digest", MODEL, DIGEST_PROMPT,
            messages, response_text=digest)
        db.set_material_digest(material_id, digest)
        db.log_event(m["user_id"], "material_digested",
                     {"material_id": material_id,
                      "chars_read": len(body), "truncated": truncated,
                      "llm_call_id": call_id},
                     source="materials")
        return digest
    except Exception as e:
        print(f"[MATERIALS] ⚠️ digest failed for {material_id}: {e}",
              flush=True)
        return None
