"""
Learning materials + magic-link tokens (offer-loop arc, PR 1).

A material is one thing the user studies from — an uploaded file, a
shared link, or a source only named in conversation. The Theo-led
walkthrough fills user_description/wants with the user's OWN words,
and reaches 'validated' only when the coach's sample (an
insider-plausible question, a next-piece cut) was confirmed by the
user. has_validated_material() is the mechanical fill condition the
material_walkthrough onboarding field will key on.

Run: ./venv/bin/python test_materials.py  (sqlite)
"""

import json
import os
import tempfile

os.environ.pop("DATABASE_URL", None)
os.environ["TZ_OFFSET_HOURS"] = "0"

import db  # noqa: E402

db.DB_PATH = os.path.join(tempfile.mkdtemp(), "test_materials.db")
db.init_db()

U = "m1"
PASS = []


def check(name, cond):
    print(f"  {'✅' if cond else '❌'} {name}")
    PASS.append(bool(cond))


def events_of(kind, user=U):
    return [r for r in db.get_events(user, limit=300) if r["kind"] == kind]


# ── 1. the three kinds ───────────────────────────────────────────────
print("1) kinds")
db.ensure_user_profile_row(U)
f_id = db.add_user_material(U, "file", title="파생상품_정리.docx",
                            orig_filename="파생상품_정리.docx",
                            extracted_text="제1장 스왑...")
l_id = db.add_user_material(U, "link", title="Karpathy — zero to hero",
                            source_url="https://youtube.com/watch?v=abc")
n_id = db.add_user_material(U, "named", title="그 책")
check("three kinds registered, ids distinct",
      len({f_id, l_id, n_id}) == 3)
check("registration events carry kind + title",
      len(events_of("material_added")) == 3
      and json.loads(events_of("material_added")[-1]["payload"])["kind"]
      in ("file", "link", "named"))
try:
    db.add_user_material(U, "telepathy")
    check("unknown kind rejected", False)
except ValueError:
    check("unknown kind rejected", True)

rows = db.get_user_materials(U)
check("listed newest first, wants decoded",
      rows[0]["id"] == n_id and rows[0]["wants"] == []
      and rows[-1]["orig_filename"] == "파생상품_정리.docx")

# ── 2. digest ────────────────────────────────────────────────────────
print("2) digest")
db.set_material_digest(f_id, "3개 장: 스왑/옵션/규제. 항목 41개.",
                       extracted_text="제1장 스왑... (전문)")
m = db.get_material(f_id)
check("digest + extracted text stored",
      "41개" in m["digest"] and "전문" in m["extracted_text"])

# ── 3. the walkthrough lands across turns ────────────────────────────
print("3) walkthrough")
db.update_material_walkthrough(f_id, user_description="업무에서 쌓은 걸 정리한 파일",
                               status="in_progress")
db.update_material_walkthrough(
    f_id,
    wants=[{"quote": "클라이언트가 물으면 바로 대답해야 해",
            "meaning": "instant recall under questioning"}])
m = db.get_material(f_id)
check("partial updates accumulate (description kept after wants-only update)",
      m["user_description"] == "업무에서 쌓은 걸 정리한 파일"
      and m["wants"][0]["quote"].startswith("클라이언트")
      and m["walkthrough_status"] == "in_progress")
check("each landing is an event",
      len(events_of("material_walkthrough_updated")) == 2)
try:
    db.update_material_walkthrough(f_id, status="perfect")
    check("unknown status rejected", False)
except ValueError:
    check("unknown status rejected", True)

# ── 4. validation is the gate ────────────────────────────────────────
print("4) validation gate")
check("not validated yet → field would stay missing",
      not db.has_validated_material(U))
db.update_material_walkthrough(f_id, status="validated")
check("coach sample confirmed → gate opens",
      db.has_validated_material(U))
check("another user is unaffected",
      not db.has_validated_material("someone_else"))

# ── 5. magic-link tokens ─────────────────────────────────────────────
print("5) tokens")
t1 = db.ensure_user_token(U)
check("token created once and stable",
      t1 == db.ensure_user_token(U) and len(t1) >= 32)
check("token → user (the /my auth check)",
      db.get_user_id_by_token(t1) == U
      and db.get_user_id_by_token("wrong") is None
      and db.get_user_id_by_token("") is None)
t2 = db.regenerate_user_token(U)
check("regeneration invalidates the leaked link",
      t2 != t1 and db.get_user_id_by_token(t1) is None
      and db.get_user_id_by_token(t2) == U)
check("regenerating a tokenless user still yields a working token",
      db.get_user_id_by_token(db.regenerate_user_token("fresh")) == "fresh")

# ── 6. upload path (materials.py) ────────────────────────────────────
print("6) upload handling")
import io  # noqa: E402

import materials  # noqa: E402
from docx import Document  # noqa: E402

buf = io.BytesIO()
_doc = Document()
_doc.add_paragraph("제1장 스왑 — 총수익스왑의 구조")
_doc.add_paragraph("조기상환 트리거: 발동 시 정산 기준일은...")
_doc.save(buf)
DOCX = buf.getvalue()

mid, err = materials.register_upload("m2", "정리본.docx", DOCX)
check("docx upload extracts real text and registers",
      err == "" and mid is not None
      and "조기상환" in db.get_material(mid)["extracted_text"])
check("original bytes are not kept anywhere",
      "orig_blob" not in db.get_material(mid))

_bad = [materials.register_upload("m2", "notes.txt", b"x")[1],
        materials.register_upload("m2", "big.pdf",
                                  b"x" * (materials.MAX_FILE_BYTES + 1))[1],
        materials.register_upload("m2", "img.pdf", b"%PDF-1.4 broken")[1],
        materials.register_upload("m2", "empty.docx", b"")[1]]
check("txt/oversize/corrupt/empty all refused with reasons",
      all(_bad) and "20MB" in _bad[1])

_lid, err = materials.register_link("m2", "https://youtube.com/watch?v=x",
                                    "Karpathy zero-to-hero")
check("link registered; junk link refused",
      err == "" and db.get_material(_lid)["kind"] == "link"
      and materials.register_link("m2", "not a url")[1] != "")

# ── 7. the one-time read ─────────────────────────────────────────────
print("7) digest")


class FakeAnthropic:
    seen = []

    def __init__(self, *a, **kw):
        pass

    class messages:
        @staticmethod
        def create(**kwargs):
            FakeAnthropic.seen.append(kwargs)

            class _B:
                type = "text"
                text = "Working notes: 2 sections, ~2 recall items."

            class _R:
                content = [_B()]
            return _R()


materials.anthropic.Anthropic = FakeAnthropic
digest = materials.digest_material(mid)
check("digest stored, event + flight record emitted",
      "recall items" in (digest or "")
      and "recall items" in db.get_material(mid)["digest"]
      and len(events_of("material_digested", "m2")) == 1)
check("the read saw the document text under the coach-notes prompt",
      "조기상환" in FakeAnthropic.seen[0]["messages"][0]["content"]
      and "coach's working notes"
      in FakeAnthropic.seen[0]["system"])
check("digesting a link (no text) is a quiet no-op",
      materials.digest_material(_lid) is None)

# ── 8. Theo emails the /my link (the real delivery path) ─────────────
print("8) link email")
import emailer  # noqa: E402


class FakeHTTP:
    sent = []

    def __init__(self, req, timeout=None):
        FakeHTTP.sent.append(req)

    def __enter__(self):
        class R:
            @staticmethod
            def read():
                return b'{"id": "re_123"}'
        return R()

    def __exit__(self, *a):
        return False


emailer.urllib.request.urlopen = FakeHTTP
os.environ.pop("RESEND_API_KEY", None)
ok, detail = emailer.send_my_link("m3", "spouse@example.com")
check("without an API key the send fails loudly, with the failure evented",
      not ok and "RESEND_API_KEY" in detail
      and len(events_of("my_link_email_failed", "m3")) == 1)

os.environ["RESEND_API_KEY"] = "re_test"
ok, detail = emailer.send_my_link("m3", "spouse@example.com")
sent_body = json.loads(FakeHTTP.sent[-1].data.decode())
check("the email carries the magic link, nothing else exotic",
      ok and detail == "re_123"
      and db.ensure_user_token("m3") in sent_body["text"]
      and sent_body["to"] == ["spouse@example.com"])
check("success records the event + stores the address on the profile",
      len(events_of("my_link_emailed", "m3")) == 1
      and (db.get_user_profile_by_id("m3") or {}).get("email")
      == "spouse@example.com")

# ── 9. the upload follow-up: the coach reaches out NOW ───────────────
print("9) material-ready follow-up")
import sms  # noqa: E402

os.environ["MESSAGING_CHANNEL"] = "sms"       # no window gate in the way
os.environ["TUTOR_USER_ID"] = "m4"
os.environ["TUTOR_USER_PHONE"] = "+15550009999"
os.environ.pop("TWILIO_ACCOUNT_SID", None)    # send becomes a logged no-op


class FakeAll:
    """sms/materials/analyze share ONE anthropic module object, so a
    single dispatching fake (the project convention) branches on the
    call's shape."""

    def __init__(self, *a, **kw):
        pass

    class messages:
        @staticmethod
        def create(**kwargs):
            if kwargs.get("tools"):
                class _B:
                    type = "tool_use"
                    input = {"step_completed": "not_applicable",
                             "step_reason": "-"}
            elif "reading a document" in kwargs.get("system", ""):
                class _B:
                    type = "text"
                    text = "Working notes: 2 sections, ~2 recall items."
            else:
                class _B:
                    type = "text"
                    text = ("정리본 읽었어 — 3장 조기상환 표가 압권이던데.\n"
                            "시간 될 때 같이 훑어보자.\n"
                            "[STEP: connect@1, spark_curiosity@1]\n"
                            "[EXPECT: reply]")

            class _R:
                content = [_B()]
            return _R()


sms.anthropic.Anthropic = FakeAll

M4 = "m4"
db.ensure_user_profile_row(M4)
db.set_agreed_goal(M4, "g"); db.save_learning_path(M4, "d", "p", "c")
db.set_ignition_marker(M4, "opens the file")
_fid, _ = materials.register_upload(M4, "정리본.docx", DOCX)
materials.digest_material(_fid)


def m4_events(kind):
    return [r for r in db.get_events(M4, limit=100) if r["kind"] == kind]


check("digest completion fires the coach's follow-up on the thread",
      len(m4_events("sms_out")) == 1
      and json.loads(m4_events("sms_out")[0]["payload"])["trigger"]
      == "material_ready"
      and "같이 훑어보자" in db.get_recent_sms_messages(M4, limit=1)
      [0]["content"])

M5 = "m5"                    # onboarding-complete user: never fires
os.environ["TUTOR_USER_ID"] = "m5"
db.ensure_user_profile_row(M5)
db.check_and_complete_onboarding(M5, force=True)
_fid2, _ = materials.register_upload(M5, "정리본.docx", DOCX)
materials.digest_material(_fid2)
check("a completed user's upload stays quiet (no focus to serve)",
      not [r for r in db.get_events(M5, limit=50)
           if r["kind"] == "sms_out"])

M6 = "m6"                    # goal not yet agreed → focus isn't walkthrough
os.environ["TUTOR_USER_ID"] = "m6"
db.ensure_user_profile_row(M6)
_fid3, _ = materials.register_upload(M6, "정리본.docx", DOCX)
materials.digest_material(_fid3)
check("upload before the goal is agreed stays quiet (focus is the goal)",
      not [r for r in db.get_events(M6, limit=50)
           if r["kind"] == "sms_out"])

os.environ["MESSAGING_CHANNEL"] = "whatsapp"  # closed window → refusal logged
os.environ["TUTOR_USER_ID"] = "m7"
M7 = "m7"
db.ensure_user_profile_row(M7)
db.set_agreed_goal(M7, "g"); db.save_learning_path(M7, "d", "p", "c")
db.set_ignition_marker(M7, "m")
_fid4, _ = materials.register_upload(M7, "정리본.docx", DOCX)
materials.digest_material(_fid4)
check("closed WhatsApp window → refusal recorded, nothing sent",
      [r for r in db.get_events(M7, limit=50)
       if r["kind"] == "material_ready_not_sent"]
      and not [r for r in db.get_events(M7, limit=50)
               if r["kind"] == "sms_out"])
os.environ["MESSAGING_CHANNEL"] = "sms"

print(f"\n{sum(PASS)}/{len(PASS)} passed")
raise SystemExit(0 if all(PASS) else 1)
